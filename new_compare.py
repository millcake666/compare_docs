import difflib
from docx import Document

document1 = Document('data/Договор поставки Товара (ООО ТОЧИНВЕСТ-ШЗМК - покупатель) 2022.docx')
document2 = Document('data/Договор поставки Товара (ООО ТОЧИНВЕСТ-ШЗМК - покупатель) 2022 новый.docx')

# Extract text from the first document
text1 = "\n".join([para.text for para in document1.paragraphs])

# Extract text from the second document
text2 = "\n".join([para.text for para in document2.paragraphs])

# Compare the text content of the two documents
diff = difflib.ndiff(text1.splitlines(), text2.splitlines())
delta = '\n'.join(x[2:] for x in diff if x.startswith('- '))

# print(delta)
